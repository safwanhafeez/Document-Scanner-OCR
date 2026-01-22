import os
import requests
import json
import base64
import tempfile
import subprocess
import re
import logging
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format='%(message)s'
)
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

app = Flask(__name__)
CORS(app)

API_KEY = os.getenv('GEMINI_API_KEY', '')

def getGeminiModel():
    modelName = 'gemini-1.5-flash-latest'
    try:
        logger.info('Discovering available Gemini models...')
        modelsUrl = f'https://generativelanguage.googleapis.com/v1beta/models?key={API_KEY}'
        response = requests.get(modelsUrl, timeout=10)
        if response.status_code == 200:
            data = response.json()
            for model in data.get('models', []):
                methods = model.get('supportedGenerationMethods', [])
                if 'generateContent' in methods:
                    name = model['name'].replace('models/', '')
                    if 'flash' in name:
                        modelName = name
                        logger.info(f'Using model: {modelName}')
                        return modelName
                    if 'pro' in name:
                        modelName = name
    except Exception as e:
        logger.error(f'Model discovery failed: {e}')
    
    logger.info(f'Using fallback model: {modelName}')
    return modelName

def getGeminiResponse(imageData):
    try:
        logger.info('Encoding image to base64...')
        base64Image = base64.b64encode(imageData).decode('utf-8')
        modelName = getGeminiModel()
        
        url = f'https://generativelanguage.googleapis.com/v1beta/models/{modelName}:generateContent?key={API_KEY}'
        headers = {'Content-Type': 'application/json'}
        
        promptText = """You are an expert OCR and technical diagram transcription system.

TASK:
1. Transcribe all text from the image accurately.
2. If you see any diagrams, charts, graphs, or technical illustrations:
   - DO NOT describe them in text.
   - Instead, write a Python script using matplotlib to RECREATE that diagram exactly.
   - Place the Python code inside these specific tags: [[DIAGRAM_CODE_START]] ... [[DIAGRAM_CODE_END]]
   - The Python code MUST save the figure to a file named 'generated_diagram.png' and close the plot.
   - Example code structure:
     import matplotlib.pyplot as plt
     fig, ax = plt.subplots()
     plt.savefig('generated_diagram.png')
     plt.close()
   - Use ONLY matplotlib and numpy.

FORMATTING:
- Output the text normally.
- Insert the diagram code blocks in the natural flow where the diagrams appear in the document."""
        
        payload = {
            'contents': [{
                'parts': [
                    {'text': promptText},
                    {
                        'inline_data': {
                            'mime_type': 'image/jpeg',
                            'data': base64Image
                        }
                    }
                ]
            }]
        }
        
        logger.info('Sending request to Gemini API...')
        response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=120)
        
        if response.status_code != 200:
            logger.error(f'Gemini API error: {response.status_code}')
            return None
            
        result = response.json()
        if 'candidates' in result and result['candidates']:
            responseText = result['candidates'][0]['content']['parts'][0]['text']
            logger.info('Successfully received response from Gemini API')
            return responseText
        else:
            logger.error('No response candidates from Gemini API')
            return None
            
    except Exception as e:
        logger.error(f'Gemini API call failed: {e}', exc_info=True)
        return None

def executeDiagramCode(code, uniqueId, workDir):
    try:
        scriptPath = os.path.join(workDir, f'diagram_{uniqueId}.py')
        imageOutputPath = os.path.join(workDir, f'diagram_{uniqueId}.png')
        
        escapedPath = imageOutputPath.replace('\\', '/')
        modifiedCode = code.replace('generated_diagram.png', escapedPath)
        modifiedCode = modifiedCode.replace(r'\implies', r'\Rightarrow')
        
        modifiedCode = modifiedCode.replace('ax.xlim(', 'ax.set_xlim(')
        modifiedCode = modifiedCode.replace('ax.ylim(', 'ax.set_ylim(')
        modifiedCode = modifiedCode.replace('ax_linear.xlim(', 'ax_linear.set_xlim(')
        modifiedCode = modifiedCode.replace('ax_linear.ylim(', 'ax_linear.set_ylim(')
        modifiedCode = modifiedCode.replace('ax_log.xlim(', 'ax_log.set_xlim(')
        modifiedCode = modifiedCode.replace('ax_log.ylim(', 'ax_log.set_ylim(')
        
        with open(scriptPath, 'w', encoding='utf-8') as f:
            f.write(modifiedCode)
        
        logger.info(f'Executing diagram generation script...')
        result = subprocess.run(
            ['python', scriptPath],
            capture_output=True,
            text=True,
            timeout=30
        )
        
        if result.returncode == 0 and os.path.exists(imageOutputPath):
            logger.info('Diagram generated successfully')
            return imageOutputPath
        else:
            logger.error(f'Diagram generation failed: {result.stderr}')
            return None
            
    except Exception as e:
        logger.error(f'Error executing diagram code: {e}', exc_info=True)
        return None

def processImageToDoc(imageData, filename, doc, workDir):
    try:
        logger.info(f'Starting conversion for: {filename}')
        
        responseText = getGeminiResponse(imageData)
        if not responseText:
            logger.error('Failed to get response from Gemini API')
            return False
        
        logger.info('Parsing response and building document...')
        doc.add_heading(f'Source: {filename}', level=1)
        
        parts = re.split(
            r'\[\[DIAGRAM_CODE_START\]\](.*?)\[\[DIAGRAM_CODE_END\]\]',
            responseText,
            flags=re.DOTALL
        )
        
        diagramCounter = 0
        for i, part in enumerate(parts):
            if i % 2 == 0:
                if part.strip():
                    doc.add_paragraph(part.strip())
            else:
                codeBlock = part.strip()
                codeBlock = codeBlock.replace('```python', '').replace('```', '')
                
                uniqueId = f'{filename}_{diagramCounter}'
                uniqueId = ''.join([c for c in uniqueId if c.isalnum() or c == '_'])
                diagramCounter += 1
                
                logger.info(f'Processing diagram {diagramCounter}...')
                imgPath = executeDiagramCode(codeBlock, uniqueId, workDir)
                
                if imgPath and os.path.exists(imgPath):
                    try:
                        doc.add_picture(imgPath, width=Inches(5))
                        logger.info(f'Diagram {diagramCounter} added to document')
                    except Exception as e:
                        logger.error(f'Failed to add diagram to document: {e}')
                        doc.add_paragraph('[Diagram Generation Failed - Image Error]')
                else:
                    doc.add_paragraph('[Diagram Generation Failed]')
        
        doc.add_page_break()
        logger.info(f'Conversion completed successfully')
        return True
        
    except Exception as e:
        logger.error(f'Error processing image {filename}: {e}', exc_info=True)
        return False

@app.route('/api/health', methods=['GET'])
def health():
    apiKeyConfigured = bool(API_KEY)
    apiWorking = False
    
    if apiKeyConfigured:
        try:
            logger.info('Testing Gemini API connection...')
            modelsUrl = f'https://generativelanguage.googleapis.com/v1beta/models?key={API_KEY}'
            response = requests.get(modelsUrl, timeout=10)
            apiWorking = response.status_code == 200
            if apiWorking:
                logger.info('Gemini API is working')
            else:
                logger.error(f'Gemini API test failed: {response.status_code}')
        except Exception as e:
            logger.error(f'Gemini API test error: {e}')
    
    return jsonify({
        'status': 'ok',
        'apiKeyConfigured': apiKeyConfigured,
        'apiWorking': apiWorking
    })

@app.route('/api/convert', methods=['POST'])
def convert():
    outputPath = None
    try:
        if 'file' not in request.files:
            logger.error('No file in request')
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            logger.error('Empty filename')
            return jsonify({'error': 'Empty filename'}), 400
        
        if not API_KEY:
            logger.error('API key not configured')
            return jsonify({'error': 'API key not configured'}), 500
        
        logger.info(f'Received conversion request for: {file.filename}')
        
        imageData = file.read()
        logger.info(f'Image size: {len(imageData)} bytes')
        
        with tempfile.TemporaryDirectory() as workDir:
            doc = Document()
            doc.add_heading('Smart OCR Conversion', level=0)
            
            success = processImageToDoc(imageData, file.filename, doc, workDir)
            
            if not success:
                logger.error('Conversion process failed')
                return jsonify({'error': 'Processing failed'}), 500
            
            outputPath = os.path.join(tempfile.gettempdir(), f'converted_{file.filename.rsplit(".", 1)[0]}.docx')
            doc.save(outputPath)
            logger.info('Document saved successfully')
            
            with open(outputPath, 'rb') as f:
                docxData = f.read()
            
            try:
                os.remove(outputPath)
            except:
                pass
            
            from io import BytesIO
            return send_file(
                BytesIO(docxData),
                as_attachment=True,
                download_name=f'converted_{file.filename.rsplit(".", 1)[0]}.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
    except Exception as e:
        logger.error(f'Conversion endpoint error: {e}', exc_info=True)
        if outputPath and os.path.exists(outputPath):
            try:
                os.remove(outputPath)
            except:
                pass
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    print('=' * 60)
    print('Smart OCR Converter Backend')
    print('=' * 60)
    if not API_KEY:
        print('WARNING: GEMINI_API_KEY environment variable not set')
        print('Please set it before making conversion requests')
    else:
        print('API Key: Configured')
    print('Server starting on http://localhost:5000')
    print('=' * 60)
    app.run(debug=False, host='0.0.0.0', port=5000)
